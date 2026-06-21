import io
import json
import os

import anthropic
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt
from flask import Flask, jsonify, request, send_file, send_from_directory
from flask_cors import CORS

from rag_engine import RAGEngine
from text_stats import TOOLS, TextStats, run_tool

FRONTEND_DIR = os.path.join(os.path.dirname(__file__), "..", "frontend")

app = Flask(__name__, static_folder=FRONTEND_DIR, static_url_path="")
CORS(app, origins="*")

rag_engine = RAGEngine()
text_stats = TextStats(rag_engine)
MAX_TOOL_ROUNDS = 4

SYSTEM_PROMPT = """אתה בלעם — סוכן ידע לפרשת שבוע, מבוסס על טקסט התורה והפרשנים הקלאסיים.
המשתמשים בך הם בעיקר רבנים ואנשי תורה שמתעניינים בפלפול ובדקויות — לא רק בשליפת ציטוטים. תפקידך להיות חד, בקיא ומפולפל, לא רק "מנוע חיפוש" בתוך המקורות שסופקו.

חוקי יסוד:
- כל טענה עובדתית על תוכן הפסוק או דברי מפרש מסוים — תתבסס על "מקורות" שסופקו לך כאן, ותצוטט במפורש: (פרק X פסוק Y — שם המפרש), ולפסוקי תורה עצמם: (פרק X פסוק Y — טקסט התורה). אל תייחס ציטוט למקור שלא הובא לך.
- מותר ורצוי להשתמש בידע תורני כללי שלך — כדי לחשוב, להעיר, להשוות בין מפרשים, להצביע על קשיים, השמטות, סתירות או דיוקי לשון, ולענות על שאלות פלפול שדורשות הבנה כללית של הפרשה ולא רק חיפוש מילולי. כשאתה עושה זאת, הבחן בבירור בין "כך כתוב במקור X" לבין הערה/פלפול עצמאי שלך (למשל: "יש להעיר ש...", "מבחינה פרשנית אפשר לשאול...").
- אמור "לא מצאתי מידע על כך במאגר" רק כשבאמת אינך יודע את התשובה — לא כתחליף למחשבה. אם אתה יודע את התשובה (גם אם היא לא כתובה במפורש באף אחד מהמקורות שסופקו), ענה אותה, וציין שזו ידיעה כללית ולא ציטוט ממקור.
- לכל שאלה כמותית על טקסט הפרשה עצמו (כמה פעמים מופיעה מילה/שורש, כמה פסוקים/מילים יש בפרשה וכו') — חובה להפעיל את הכלי המתאים (count_word_in_parasha / get_parasha_stats) ולהתבסס על תוצאתו המדויקת. אל תנחש ואל תסתמך על הערכה.
- עברית בלבד. תשובה חדה וממוקדת בעיקר הקושי או הפלפול — לא תשובה גנרית.

מקורות שנמצאו לשאלה הנוכחית:
{context}"""


def _format_source_label(meta):
    chapter = meta.get("chapter")
    verse = meta.get("verse")
    commentator = meta.get("commentator_name")
    if commentator:
        return f"פרק {chapter} פסוק {verse} — {commentator}"
    return f"פרק {chapter} פסוק {verse} — טקסט התורה"


def retrieve_context(query, n=6):
    results = rag_engine.search(query, n=n)
    documents = results["documents"][0]
    metadatas = results["metadatas"][0]

    context_lines = []
    sources = []
    for text, meta in zip(documents, metadatas):
        label = _format_source_label(meta)
        context_lines.append(f"[{label}]\n{text}")
        sources.append({
            "chapter": meta.get("chapter"),
            "verse": meta.get("verse"),
            "ref_he": meta.get("ref_he"),
            "commentator_name": meta.get("commentator_name"),
            "source_type": meta.get("source_type"),
            "source_url": meta.get("source_url"),
        })
    return "\n\n".join(context_lines), sources


def last_user_message(messages):
    for msg in reversed(messages):
        if msg.get("role") == "user":
            content = msg.get("content")
            if isinstance(content, list):
                return " ".join(b.get("text", "") for b in content if isinstance(b, dict))
            return content
    return ""


@app.route("/", methods=["GET"])
def index():
    return send_from_directory(app.static_folder, "index.html")


@app.route("/health", methods=["GET"])
def health():
    return jsonify({"status": "ok", "doc_count": rag_engine.count()})


@app.route("/chat", methods=["POST", "OPTIONS"])
def chat():
    if request.method == "OPTIONS":
        return "", 204

    data = request.json or {}
    messages = data.get("messages", [])
    api_key = data.get("api_key")

    if not api_key:
        return jsonify({"error": "חובה להזין מפתח Anthropic API"}), 400
    if not messages:
        return jsonify({"error": "לא התקבלה שאלה"}), 400

    query = last_user_message(messages)
    context, sources = retrieve_context(query)

    if not context:
        context = "(לא נמצאו מקורות רלוונטיים במאגר עבור שאלה זו)"

    system = SYSTEM_PROMPT.format(context=context)

    try:
        client = anthropic.Anthropic(api_key=api_key)
        conversation = list(messages)
        response = None
        for _ in range(MAX_TOOL_ROUNDS):
            response = client.messages.create(
                model="claude-sonnet-4-6",
                max_tokens=2000,
                system=system,
                tools=TOOLS,
                messages=conversation,
            )
            if response.stop_reason != "tool_use":
                break

            conversation.append({"role": "assistant", "content": response.content})
            tool_results = []
            for block in response.content:
                if block.type == "tool_use":
                    result = run_tool(text_stats, block.name, block.input)
                    tool_results.append({
                        "type": "tool_result",
                        "tool_use_id": block.id,
                        "content": json.dumps(result, ensure_ascii=False),
                    })
            conversation.append({"role": "user", "content": tool_results})

        text_blocks = [b.text for b in response.content if hasattr(b, "text")]
        reply = text_blocks[0] if text_blocks else ""
    except anthropic.AuthenticationError:
        return jsonify({"error": "מפתח ה-API שגוי או לא תקף"}), 401
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500
    finally:
        api_key = None  # never retained past this request

    return jsonify({"reply": reply, "sources": sources})


def _set_rtl(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.makeelement(qn("w:bidi"), {})
    p_pr.append(bidi)


def build_docx(question, context_groups):
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(12)

    title = doc.add_heading(question, level=1)
    _set_rtl(title)

    for group in context_groups:
        verse_p = doc.add_paragraph()
        verse_run = verse_p.add_run(f"{group['ref_he']}: {group['verse_text_hebrew']}")
        verse_run.bold = True
        _set_rtl(verse_p)

        for commentary in group["commentaries"]:
            c_p = doc.add_paragraph()
            c_run = c_p.add_run(f"{commentary['commentator_name']}: {commentary['text']}")
            _set_rtl(c_p)
            c_p.paragraph_format.left_indent = Pt(18)

        doc.add_paragraph()

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


@app.route("/generate-docx", methods=["POST", "OPTIONS"])
def generate_docx():
    if request.method == "OPTIONS":
        return "", 204

    data = request.json or {}
    messages = data.get("messages", [])
    if not messages:
        return jsonify({"error": "לא התקבלה שאלה"}), 400

    query = last_user_message(messages)
    results = rag_engine.search(query, n=10)
    documents = results["documents"][0]
    metadatas = results["metadatas"][0]

    groups = {}
    for text, meta in zip(documents, metadatas):
        key = (meta.get("chapter"), meta.get("verse"))
        if key not in groups:
            groups[key] = {
                "ref_he": meta.get("ref_he"),
                "verse_text_hebrew": meta.get("verse_text_hebrew"),
                "commentaries": [],
            }
        if meta.get("source_type") == "commentary":
            groups[key]["commentaries"].append({
                "commentator_name": meta.get("commentator_name"),
                "text": text,
            })

    ordered_groups = [groups[k] for k in sorted(groups.keys())]

    if not ordered_groups:
        return jsonify({"error": "לא נמצאו מקורות רלוונטיים במאגר עבור שאלה זו"}), 404

    buffer = build_docx(query, ordered_groups)
    return send_file(
        buffer,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name="bilam.docx",
    )


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5005))
    app.run(host="0.0.0.0", port=port)
